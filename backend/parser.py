from pathlib import Path


def _read_pdf(file_path: Path) -> str:
    """Read all text from a PDF file."""
    from pypdf import PdfReader
    reader = PdfReader(str(file_path))
    parts = []
    for i, page in enumerate(reader.pages):
        text = page.extract_text()
        if text:
            parts.append(f"--- 第 {i + 1} 页 ---\n{text.strip()}")
    return "\n\n".join(parts)


def _read_docx(file_path: Path) -> str:
    """Read all text from a DOCX file, including tables."""
    from docx import Document
    doc = Document(str(file_path))
    parts = []
    for para in doc.paragraphs:
        if para.text.strip():
            parts.append(para.text.strip())
    for i, table in enumerate(doc.tables):
        parts.append(f"\n[表格 {i + 1}]")
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells]
            parts.append(" | ".join(cells))
    return "\n".join(parts)


def parse_resume(file_path: Path) -> str:
    """Parse a resume file (PDF or DOCX) and return plain text."""
    ext = file_path.suffix.lower()
    if ext == ".pdf":
        return _read_pdf(file_path)
    elif ext in (".docx", ".doc"):
        return _read_docx(file_path)
    else:
        raise ValueError(f"Unsupported file type: {ext}")
