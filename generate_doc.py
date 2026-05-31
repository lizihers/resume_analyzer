from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = Document()

style = doc.styles['Normal']
font = style.font
font.name = 'Calibri'
font.size = Pt(11)

# 封面
doc.add_paragraph()
doc.add_paragraph()
title = doc.add_paragraph()
title.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = title.add_run('AI 简历分析器 + 岗位匹配')
run.font.size = Pt(28)
run.font.bold = True
run.font.color.rgb = RGBColor(0x6c, 0x8c, 0xff)

doc.add_paragraph()
subtitle = doc.add_paragraph()
subtitle.alignment = WD_ALIGN_PARAGRAPH.CENTER
run = subtitle.add_run('使用教程')
run.font.size = Pt(18)
run.font.color.rgb = RGBColor(0x8b, 0x8f, 0xa3)

doc.add_paragraph()
info = doc.add_paragraph()
info.alignment = WD_ALIGN_PARAGRAPH.CENTER
info.add_run('版本 1.0  |  2026年5月').font.size = Pt(11)

doc.add_page_break()

# 目录
doc.add_heading('目录', level=1)
toc = [
    '1. 系统概述',
    '2. 快速开始',
    '3. 注册与登录',
    '4. 简历分析与学历评估',
    '5. 岗位匹配',
    '6. 职位推荐',
    '7. 历史记录',
    '8. 隐私保护说明',
    '9. 公网部署',
    '10. 常见问题',
]
for item in toc:
    doc.add_paragraph(item)
doc.add_page_break()

# 1
doc.add_heading('1. 系统概述', level=1)
doc.add_paragraph(
    'AI 简历分析器是一款全栈 AI 工具，帮助求职者深度分析简历质量、'
    '评估学历竞争力、匹配岗位需求、推荐目标公司。系统采用前后端分离架构，'
    '支持多人注册使用，每个人的数据完全隔离。'
)
doc.add_paragraph('核心功能：')
features = [
    '简历解析：上传 PDF/DOCX 格式简历，自动提取文本',
    '隐私脱敏：自动隐藏姓名、电话、邮箱等个人信息，保护用户隐私',
    'AI 深度分析：提取技能、经历、教育背景，四维度评分（完整度/影响力/关键词/综合）',
    '学历竞争力：评估学校档次、学历层次、专业匹配度、综合竞争力',
    '岗位匹配：粘贴 JD 或输入招聘链接，AI 对比简历与岗位要求',
    '简历优化方向：关键词加入建议、经历改写建议、优先学习技能',
    '职位推荐：AI 推荐 3-4 个岗位方向（含预估薪资），匹配 6-8 家目标公司',
    '100+ 公司招聘官网：互联网/芯片/AI/金融/汽车/快消/医药等全行业覆盖，一键直达',
    '历史记录：保存分析记录，支持随时查看和删除',
]
for f in features:
    doc.add_paragraph(f, style='List Bullet')
doc.add_page_break()

# 2
doc.add_heading('2. 快速开始', level=1)

doc.add_heading('2.1 本地启动', level=2)
doc.add_paragraph('打开终端（PowerShell 或 CMD），执行以下命令：')
doc.add_paragraph('cd C:\\Users\\ZWY\\resume_analyzer')
doc.add_paragraph('python -m uvicorn backend.main:app --host 127.0.0.1 --port 8000')
doc.add_paragraph('看到 "Application startup complete" 后，打开浏览器访问：')
doc.add_paragraph('http://127.0.0.1:8000')

doc.add_heading('2.2 一键启动（后台运行）', level=2)
doc.add_paragraph('双击项目目录下的 deploy.bat 文件，自动启动服务和公网隧道。')

doc.add_heading('2.3 公开分享（临时）', level=2)
doc.add_paragraph('在另一个终端运行：')
doc.add_paragraph('cloudflared tunnel --url http://localhost:8000 --protocol http2 --no-autoupdate')
doc.add_paragraph('将生成的 https://xxx.trycloudflare.com 地址发给他人即可访问。\n注意：你的电脑需要保持开机和联网。')

doc.add_heading('2.4 永久部署', level=2)
doc.add_paragraph('参考第9章「公网部署」将服务部署到云平台，实现永久在线。')
doc.add_page_break()

# 3
doc.add_heading('3. 注册与登录', level=1)
doc.add_paragraph('首次使用需要注册账号。每个用户的数据完全隔离，互不可见。')

doc.add_heading('3.1 注册', level=2)
doc.add_paragraph('① 打开首页，点击「注册」标签')
doc.add_paragraph('② 输入用户名（2-30 个字符）和密码（至少 4 个字符）')
doc.add_paragraph('③ 点击「注册」按钮，成功后自动登录进入主界面')

doc.add_heading('3.2 登录', level=2)
doc.add_paragraph('① 打开首页，确认选中「登录」标签')
doc.add_paragraph('② 输入已注册的用户名和密码')
doc.add_paragraph('③ 点击「登录」按钮进入主界面')

doc.add_heading('3.3 退出登录', level=2)
doc.add_paragraph('登录后，页面右上角显示你的用户名，点击右侧的「退出」按钮即可。')

doc.add_heading('3.4 安全说明', level=2)
doc.add_paragraph(
    '密码使用 PBKDF2 加盐哈希存储，即使数据库泄露也无法还原明文密码。'
    '登录 Token 有效期为 7 天，过期后需重新登录。'
)
doc.add_page_break()

# 4
doc.add_heading('4. 简历分析与学历评估', level=1)

doc.add_heading('4.1 上传简历', level=2)
doc.add_paragraph('① 登录后默认进入「简历分析」页面')
doc.add_paragraph('② 点击上传区域选择文件，或直接拖拽 PDF/DOCX 文件到虚线框内')
doc.add_paragraph('③ 系统自动解析文件内容，显示「解析完成」及字符数（标注"已脱敏"）')
doc.add_paragraph('④ 点击「查看解析文本」可预览解析结果（个人信息已自动脱敏）')

doc.add_heading('4.2 AI 分析', level=2)
doc.add_paragraph('① 解析完成后，点击「开始 AI 分析」按钮')
doc.add_paragraph('② 等待约 10-30 秒（取决于简历长度），AI 生成分析结果')

doc.add_heading('4.3 分析结果解读', level=2)
doc.add_paragraph('分析结果包含以下板块：')
sections = [
    '学历竞争力：学校档次（985/211/双一流/省重点等，绿色=名校，黄色=中等，蓝色=普通）、学历层次、专业匹配度（0-100）、综合竞争力（0-100）、学历优势与限制分析',
    '综合评分：四维度打分——完整度/影响力/关键词优化/综合，各带进度条',
    '技能标签：技术技能（蓝色标签）和软技能（紫色标签）',
    '经历梳理：每段工作/项目经历、公司/组织、时间、亮点',
    '优势列表：简历的核心优势（绿色圆点标记）',
    '待改进：简历的薄弱环节（红色圆点标记）',
    '改进建议：按类别（content 内容/format 格式/keyword 关键词）给出具体优化建议',
]
for s in sections:
    doc.add_paragraph(s, style='List Bullet')
doc.add_page_break()

# 5
doc.add_heading('5. 岗位匹配', level=1)
doc.add_paragraph('此功能将你的简历与具体职位描述（JD）进行对比分析。')

doc.add_heading('5.1 输入职位描述', level=2)
doc.add_paragraph('方式一（推荐）：在顶部的 URL 抓取栏粘贴招聘链接，点击「抓取」按钮，系统自动提取 JD 内容')
doc.add_paragraph('方式二：直接在下方文本框粘贴职位描述文本')

doc.add_heading('5.2 开始匹配', level=2)
doc.add_paragraph('① 确保已在「简历分析」中上传并分析了简历')
doc.add_paragraph('② 输入 JD 后，点击「开始岗位匹配」按钮')
doc.add_paragraph('③ 等待 AI 分析完成')

doc.add_heading('5.3 匹配结果解读', level=2)
match_items = [
    '匹配度：综合匹配分数 0-100，绿色进度条展示',
    '匹配技能：简历中与 JD 要求一致的技能（绿色标签）',
    '缺失技能：JD 要求但简历中缺少的技能（红色标签）',
    '差距分析：能力和经验上的具体差距',
    '改进建议：针对差距的行动建议',
    '简历优化方向——关键词加入：应增加到简历中的 JD 关键词',
    '简历优化方向——经历改写建议：如何改写现有经历来匹配 JD 要求',
    '简历优化方向——优先学习技能：按优先级排列的待学习技能及理由',
]
for m in match_items:
    doc.add_paragraph(m, style='List Bullet')
doc.add_page_break()

# 6
doc.add_heading('6. 职位推荐', level=1)
doc.add_paragraph('此功能基于你的简历（技能+学历+经验），由 AI 推荐适合的岗位方向。')

doc.add_heading('6.1 使用方式', level=2)
doc.add_paragraph('① 先完成简历分析（第4章）')
doc.add_paragraph('② 切换到「职位推荐」Tab')
doc.add_paragraph('③ 点击「开始岗位推荐」按钮')

doc.add_heading('6.2 推荐结果', level=2)
rec = [
    '推荐岗位方向：3-4 个岗位名称、适合原因、预估薪资范围',
    '目标公司：6-8 家匹配公司，按推荐度排序（推荐/备选），附带匹配原因',
    '招聘官网：每条公司卡片右侧均有「招聘官网 ↗」链接，点击直接跳转到该公司的官方招聘页面',
    '职业发展建议：2-3 条职业规划建议',
]
for r in rec:
    doc.add_paragraph(r, style='List Bullet')

doc.add_heading('6.3 覆盖行业', level=2)
doc.add_paragraph(
    '系统内置 100+ 家公司招聘官网，涵盖以下行业：互联网/科技、AI/大模型、半导体/芯片、'
    '汽车/新能源、金融/银行、外资企业、制造/工业、生物医药、咨询/审计、'
    '消费品/零售、物流/运输、教育/媒体、游戏等。'
)
doc.add_page_break()

# 7
doc.add_heading('7. 历史记录', level=1)
doc.add_paragraph('系统自动保存每次 AI 分析记录，方便随时回顾。')
doc.add_paragraph('① 切换到「历史记录」Tab')
doc.add_paragraph('② 点击「查看」按钮恢复到之前的分析结果')
doc.add_paragraph('③ 点击「删除」按钮删除不需要的记录')
doc.add_paragraph(
    '注意：每条记录包含简历文本、分析结果和匹配结果。'
    '不同用户之间的记录完全隔离，只能看到自己的数据。'
)
doc.add_page_break()

# 8
doc.add_heading('8. 隐私保护说明', level=1)
doc.add_paragraph('系统从设计层面重视用户隐私保护：')
privacy = [
    '自动脱敏：简历上传后，系统自动识别并隐藏姓名、手机号、电子邮箱等个人身份信息，用星号（*）替换。例如：「张三」替换为「张*」，「13812345678」替换为「138****5678」，「zhangsan@example.com」替换为「zh****an@example.com」',
    'AI 不接触原始信息：脱敏后的文本才发送给 AI 进行分析，AI 永远看不到真实的姓名、电话、邮箱',
    '数据隔离：每个用户只能看到自己的数据，不同用户的数据通过 user_id 严格隔离',
    '密码安全：密码使用 PBKDF2 加盐哈希存储，不以明文形式保存',
    '本地解析：简历文件解析在你的本地服务器上完成，原始文件不会被上传到第三方',
    '数据库保护：SQLite 数据库文件仅存储在服务器端，用户无法直接访问',
]
for p in privacy:
    doc.add_paragraph(p, style='List Bullet')
doc.add_page_break()

# 9
doc.add_heading('9. 公网部署', level=1)
doc.add_paragraph('本地服务关掉就停了。要让服务永久在线，需要部署到云平台。')

doc.add_heading('方案 A：Cloudflare Tunnel（临时分享，免费）', level=2)
doc.add_paragraph('适用场景：临时分享给朋友、面试官查看，需要你的电脑保持开机')
doc.add_paragraph('① 安装 cloudflared（已安装则跳过）：winget install Cloudflare.cloudflared')
doc.add_paragraph('② 先启动本地服务，再启动隧道：cloudflared tunnel --url http://localhost:8000 --protocol http2 --no-autoupdate')
doc.add_paragraph('③ 将生成的 https://xxx.trycloudflare.com 地址发给他人')

doc.add_heading('方案 B：Sealos 云部署（永久在线，推荐）', level=2)
doc.add_paragraph('适用场景：需要 7×24 小时在线，不依赖自己电脑')
doc.add_paragraph('① 访问 github.com/new 创建仓库，将代码推送上去')
doc.add_paragraph('② 访问 sealos.run 注册账号')
doc.add_paragraph('③ 新建应用 → 从源代码构建 → 关联 GitHub 仓库')
doc.add_paragraph('④ 设置环境变量：OPENAI_API_KEY=你的密钥 / OPENAI_BASE_URL=https://api.deepseek.com / OPENAI_MODEL=deepseek-chat')
doc.add_paragraph('⑤ 端口填 8000，点击部署')
doc.add_paragraph('⑥ 几分钟后获得 xxx.sealos.run 域名，永久在线')

doc.add_heading('方案 C：阿里云/腾讯云轻量服务器', level=2)
doc.add_paragraph('适用场景：有预算（约 50 元/月），需要稳定和国内低延迟')
doc.add_paragraph('① 购买一台轻量应用服务器（最低配置即可）')
doc.add_paragraph('② 安装 Docker 和 Git，克隆代码')
doc.add_paragraph('③ 执行 docker build -t resume-analyzer . 构建镜像')
doc.add_paragraph('④ 执行 docker run -d -p 8000:8000 --env-file .env resume-analyzer 启动')
doc.add_paragraph('⑤ 在安全组开放 8000 端口，绑定域名（可选）')
doc.add_page_break()

# 10
doc.add_heading('10. 常见问题', level=1)
faqs = [
    ('Q: 上传简历后提示解析失败？',
     'A: 确认文件格式为 PDF 或 DOCX（不支持图片格式或扫描版 PDF）。如果是扫描版 PDF（图片型），需要先用 OCR 工具转换为文字版。'),
    ('Q: AI 分析时提示 API 密钥无效？',
     'A: 检查 .env 文件中的 OPENAI_API_KEY 是否正确。可以去 platform.deepseek.com 重新获取。'),
    ('Q: 为什么我的个人信息显示为星号？',
     'A: 这是系统的隐私保护功能，简历上传后自动脱敏姓名、电话、邮箱。AI 分析的也是脱敏后的文本，不会泄露你的个人信息。'),
    ('Q: 学历分析中为什么判定我的学校档次不准？',
     'A: 学校档次由 AI 根据学校名称公开信息推断，仅供参考。如果不准确，可以忽略该项评估，重点关注技能和经历分析。'),
    ('Q: 岗位匹配时 JD 链接抓取失败？',
     'A: 部分网站（如知乎、部分招聘网站）有反爬机制，可能无法抓取。建议直接复制粘贴 JD 文本。'),
    ('Q: 如何让朋友也使用这个工具？',
     'A: 部署到公网后（参考第9章），朋友通过浏览器访问公网地址，注册自己的账号即可使用。每个人的数据完全隔离。'),
    ('Q: 可以同时上传多份简历吗？',
     'A: 每次只能上传一份。但可以在历史记录中随时切换查看不同的分析结果。'),
    ('Q: 分析结果可以导出吗？',
     'A: 当前版本暂不支持导出功能，后续会加入 PDF/Markdown 导出功能。'),
]
for q, a in faqs:
    p = doc.add_paragraph()
    run = p.add_run(q)
    run.bold = True
    doc.add_paragraph(a)

# Save
output = 'C:/Users/ZWY/Desktop/AI简历分析器使用教程.docx'
doc.save(output)
print(f'Done: {output}')
